"""
LangChain RAG Pipeline — Trợ lý Số Cán bộ Xã
Dùng LangChain 0.3+ với:
  - OllamaEmbeddings (all-minilm)
  - OllamaLLM (qwen2.5:14b)
  - MongoDBAtlasVectorSearch
  - RecursiveCharacterTextSplitter
"""

import logging
import hashlib
import io
import time
from datetime import datetime, timezone
from typing import Optional, Generator, List, Dict, Any
import json
from langchain_ollama import OllamaEmbeddings, OllamaLLM
from langchain_mongodb import MongoDBAtlasVectorSearch
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from pymongo import MongoClient

from app.core.config import settings

logger = logging.getLogger(__name__)

# ── Streaming RAG ─────────────────────────────────────────────────────────────


def rag_chat_stream(
    query: str, session_id: str, history: list[dict] = None, department: str = "public"
) -> Generator[str, None, None]:
    """
    Streaming RAG — Trả về từng chunk JSON chuỗi cho Client.
    """
    if history is None:
        history = []

    # 1. Search (Sync - but can be async later)
    search_result = search_vectors(query, department=department)
    results = search_result["results"]

    if not results:
        yield json.dumps(
            {"type": "error", "content": "Tôi chưa tìm được thông tin."}
        ) + "\n"
        return

    # Metadata
    yield json.dumps(
        {
            "type": "metadata",
            "sources": results,
            "search_time_ms": search_result["search_time_ms"],
            "cached": search_result["cached"],
            "query_vector": search_result["query_vector"][:10],
            "query_vector_dim": len(search_result["query_vector"]),
        },
        ensure_ascii=False,
    ) + "\n"

    # Context
    context = build_context_from_results(results)
    history_str = ""
    if history:
        recent = history[-4:]
        history_str = "\n".join(
            f"{'User' if m['role']=='user' else 'Bot'}: {m['content']}" for m in recent
        )

    # 3. Stream
    llm = get_llm()
    prompt = PromptTemplate.from_template(PROMPT_TEMPLATE)
    chain = prompt | llm | StrOutputParser()

    full_answer = ""
    for chunk in chain.stream(
        {
            "context": context,
            "history": history_str or "Chưa có lịch sử hội thoại.",
            "question": query,
        }
    ):
        full_answer += chunk
        yield json.dumps({"type": "text", "content": chunk}, ensure_ascii=False) + "\n"

    # Ghi nhận lịch sử khi stream kết thúc
    history.append({"role": "user", "content": query})
    history.append({"role": "assistant", "content": full_answer})
    
    # Import save_history here avoiding circular import
    from app.main import save_history
    save_history(session_id, history)

    yield json.dumps({"type": "done"}, ensure_ascii=False) + "\n"


# ── Khởi tạo các thành phần LangChain ─────────────────────────────────────────


def get_embeddings() -> OllamaEmbeddings:
    return OllamaEmbeddings(
        model=settings.OLLAMA_EMBED_MODEL,
        base_url=settings.OLLAMA_BASE_URL,
    )


def get_llm() -> OllamaLLM:
    # Lay num_predict tu settings hoac mac dinh 150
    n_predict = getattr(settings, "OLLAMA_NUM_PREDICT", 150)
    logger.info(f"Creating LLM with num_predict={n_predict}")

    return OllamaLLM(
        model=settings.OLLAMA_LLM_MODEL,
        base_url=settings.OLLAMA_BASE_URL,
        temperature=0.1,
        num_predict=n_predict,
    )


def get_vector_store() -> MongoDBAtlasVectorSearch:
    client = MongoClient(settings.MONGO_URI)
    collection = client[settings.MONGO_DB_NAME][settings.COLLECTION_DOCUMENTS]
    return MongoDBAtlasVectorSearch(
        collection=collection,
        embedding=get_embeddings(),
        index_name=settings.VECTOR_INDEX_NAME,
        text_key="content",
        embedding_key="embedding",
        relevance_score_fn="cosine",
    )


# ── Text Splitter ──────────────────────────────────────────────────────────────


def get_splitter() -> RecursiveCharacterTextSplitter:
    """RecursiveCharacterTextSplitter: tách thông minh theo đoạn văn, câu, từ."""
    return RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        length_function=len,
    )


# ── Đọc file ───────────────────────────────────────────────────────────────────


def extract_text_from_bytes(content: bytes, file_type: str) -> list[dict]:
    """
    Trích xuất text từ file bytes.
    Trả về list[{text, page_num}]
    """
    pages = []

    if file_type == ".txt":
        text = content.decode("utf-8", errors="ignore")
        pages.append({"text": text, "page_num": 1})

    elif file_type == ".pdf":
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(content))
        for i, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append({"text": text, "page_num": i})

    elif file_type == ".docx":
        import docx

        doc = docx.Document(io.BytesIO(content))
        full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        pages.append({"text": full_text, "page_num": 1})

    return pages


# ── Ingest Pipeline ────────────────────────────────────────────────────────────


def ingest_file(
    file_type: str,
    file_id: str,
    department: str = "public",
) -> dict:
    """
    Pipeline nạp tài liệu vào MongoDB bằng LangChain.
    Tự động chunking, embedding và lưu trữ — không cần chỉ định topic.

    Returns: {chunks_total, chunks_saved, skipped}
    """
    file_hash = hashlib.sha256(content).hexdigest()

    # Đọc nội dung
    pages = extract_text_from_bytes(content, file_type)
    if not pages:
        raise ValueError("File không có nội dung văn bản.")

    # Tạo LangChain Documents từ các trang
    raw_docs = []
    for page in pages:
        doc = Document(
            page_content=page["text"],
            metadata={
                "file_id": file_id,
                "source": file_name,
                "file_name": file_name,
                "file_type": file_type.lstrip("."),
                "file_hash": file_hash,
                "page_num": page["page_num"],
                "department": department,
                "ingested_at": datetime.now(timezone.utc).isoformat(),
            },
        )
        raw_docs.append(doc)

    # Chunking bằng RecursiveCharacterTextSplitter
    splitter = get_splitter()
    chunks = splitter.split_documents(raw_docs)

    # Thêm chunk index vào metadata
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_index"] = i
        chunk.metadata["total_chunks"] = len(chunks)
        chunk.metadata["doc_id"] = f"{file_id}_chunk_{i}"

    # Embed + lưu vào MongoDB qua LangChain
    client = MongoClient(settings.MONGO_URI)
    collection = client[settings.MONGO_DB_NAME][settings.COLLECTION_DOCUMENTS]

    embeddings = get_embeddings()

    saved = 0
    skipped = 0

    # Nhúng từng chunk để theo dõi tiến độ và xử lý lỗi
    for i, chunk in enumerate(chunks):
        text = chunk.page_content
        doc_id = chunk.metadata["doc_id"]

        # Kiểm tra trùng lặp
        if collection.find_one({"metadata.doc_id": doc_id}):
            skipped += 1
            continue

        try:
            # Tạo embedding cho từng chunk
            vector = embeddings.embed_query(text)
            
            # Lưu vào MongoDB
            mongo_doc = {
                "doc_id": doc_id,
                "content": text,
                "embedding": vector,
                "metadata": chunk.metadata,
                "created_at": datetime.now(timezone.utc),
            }
            collection.insert_one(mongo_doc)
            saved += 1
        except Exception as e:
            logger.error(f"❌ Lỗi embedding chunk {i} của file {file_name}: {e}")
            logger.error(f"Nội dung lỗi: {text[:200]}...")
            continue

    logger.info(f"✅ Ingest '{file_name}': {saved} chunks saved, {skipped} skipped")
    client.close()

    return {
        "chunks_total": len(chunks),
        "chunks_saved": saved,
        "skipped": skipped,
    }


# ── Vector Search (với Cache) ──────────────────────────────────────────────────


def get_query_embedding(query: str) -> list[float]:
    """Tạo embedding từ câu hỏi và trả về vector."""
    embeddings = get_embeddings()
    vector = embeddings.embed_query(query)
    return vector


def search_vectors(query: str, top_k: int = None, department: str = None) -> dict:
    """
    H\u1ec7 th\u1ed1ng Hybrid Search: K\u1ebft h\u1ee3p Vector Search v\u00e0 Keyword Search d\u00f9ng RRF.
    Tr\u1ea3 v\u1ec1: {query_vector, results, cached, search_time_ms}
    """
    if top_k is None:
        top_k = settings.TOP_K_RESULTS

    logger.info(f"Hybrid Search for '{query[:30]}...' with top_k={top_k}")

    client = MongoClient(settings.MONGO_URI)
    db = client[settings.MONGO_DB_NAME]
    docs_col = db[settings.COLLECTION_DOCUMENTS]
    cache_col = db[settings.COLLECTION_VECTOR_CACHE]

    # Chu\u1ea9n h\u00f3a query \u0111\u1ec3 cache chính x\u00e1c
    query_norm = query.strip().lower()
    query_hash = hashlib.md5(f"{query_norm}_{department}".encode()).hexdigest()
    cached = cache_col.find_one({"query_hash": query_hash, "top_k": top_k})

    if cached:
        logger.info(f"\u26a1 Cache hit: '{query[:40]}'")
        client.close()
        return {
            "query_vector": cached["query_vector"],
            "results": cached["results"],
            "cached": True,
            "search_time_ms": 0,
        }

    t0 = time.time()

    # 1. L\u1ea5y Vector Embedding cho query
    query_vector = get_query_embedding(query)

    # Thêm filter theo bộ phận (hoặc public)
    search_filter = {}
    if department:
        search_filter = {"$or": [{"metadata.department": department}, {"metadata.department": "public"}]}

    # 2. Vector Search (Semantic)
    try:
        # Thử VectorSearch trước
        vector_results = list(
            docs_col.aggregate(
                [
                    {
                        "$vectorSearch": {
                            "index": settings.VECTOR_INDEX_NAME,
                            "path": "embedding",
                            "queryVector": query_vector,
                            "numCandidates": top_k * 10,
                            "limit": top_k * 2,
                            "filter": search_filter if search_filter else None
                        }
                    },
                    {
                        "$project": {
                            "_id": 0,
                            "content": 1,
                            "metadata": 1,
                            "vector_score": {"$meta": "vectorSearchScore"},
                        }
                    },
                ]
            )
        )
    except Exception:
        # NATIVE MATH FALLBACK: Tính Cosine Similarity bằng Aggregation Pipeline (Work trên mọi bản Mongo Local)
        # Score = Sum(Q_i * D_i) / (Norm(Q) * Norm(D))
        # Vì Norm(Q) là hằng số với 1 query nên có thể bỏ qua để tối ưu.

        # Tạo biểu thức $multiply cho top 768 dims (Nặng nhưng Native)
        # Để tối ưu, ta chỉ lấy top 100 docs bằng Keyword hoặc Filter trước, hoặc scan hết nếu dataset nhỏ.

        dot_product_expr = {
            "$sum": [
                {"$multiply": [{"$arrayElemAt": ["$embedding", i]}, query_vector[i]]}
                for i in range(len(query_vector))
            ]
        }

        vector_results = list(
            docs_col.aggregate(
                [
                    {
                        "$project": {
                            "_id": 0,
                            "content": 1,
                            "metadata": 1,
                            "vector_score": dot_product_expr,
                        }
                    },
                    {"$sort": {"vector_score": -1}},
                    {"$limit": top_k * 2},
                ]
            )
        )

    # 3. Keyword Search (Full-text) d\u00f9ng $text index
    keyword_query = {"$text": {"$search": query}}
    if department:
        keyword_query["$or"] = [{"metadata.department": department}, {"metadata.department": "public"}]

    keyword_results = list(
        docs_col.find(
            keyword_query,
            {"_id": 0, "content": 1, "metadata": 1, "score": {"$meta": "textScore"}},
        )
        .sort([("score", {"$meta": "textScore"})])
        .limit(top_k * 2)
    )

    # 4. Reciprocal Rank Fusion (RRF)
    # Thu\u1eadt to\u00e1n h\u1ee3p nh\u1ea5t k\u1ebft qu\u1ea3 chuy\u00ean nghi\u1ec7p
    k_const = 60
    rrf_scores = {}  # doc_id -> score
    docs_map = {}  # doc_id -> content/metadata

    # X\u1ebfp h\u1ea1ng t\u1eeb Vector
    for rank, doc in enumerate(vector_results, 1):
        doc_id = doc["metadata"].get("doc_id")
        rrf_scores[doc_id] = rrf_scores.get(doc_id, 0) + (1.0 / (rank + k_const))
        docs_map[doc_id] = doc
        docs_map[doc_id]["source_type"] = "vector"

    # X\u1ebfp h\u1ea1ng t\u1eeb Keyword
    for rank, doc in enumerate(keyword_results, 1):
        doc_id = doc["metadata"].get("doc_id")
        rrf_scores[doc_id] = rrf_scores.get(doc_id, 0) + (1.0 / (rank + k_const))
        if doc_id not in docs_map:
            docs_map[doc_id] = doc
            docs_map[doc_id]["source_type"] = "keyword"
        else:
            docs_map[doc_id]["source_type"] = "hybrid"

    # S\u1eafp h\u1ea1ng l\u1ea1i v\u00e0 l\u1ea5y Top K
    sorted_doc_ids = sorted(
        rrf_scores.keys(), key=lambda x: rrf_scores[x], reverse=True
    )[:top_k]

    elapsed_ms = round((time.time() - t0) * 1000, 1)

    results = []
    for doc_id in sorted_doc_ids:
        r = docs_map[doc_id]
        results.append(
            {
                "doc_id": doc_id,
                "content": r.get("content", ""),
                "content_preview": r.get("content", "")[:200],
                "source": r.get("metadata", {}).get("source", ""),
                "file_name": r.get("metadata", {}).get("file_name", ""),
                "page_num": r.get("metadata", {}).get("page_num", 1),
                "chunk_index": r.get("metadata", {}).get("chunk_index", 0),
                "score": round(rrf_scores[doc_id], 6),
                "search_type": r.get("source_type", "unknown"),
            }
        )

    # L\u01b0u v\u00e0o cache
    cache_col.insert_one(
        {
            "query_hash": query_hash,
            "query": query,
            "query_vector": query_vector,
            "top_k": top_k,
            "results": results,
            "search_time_ms": elapsed_ms,
            "created_at": datetime.now(timezone.utc),
        }
    )

    logger.info(
        f"\ud83d\udd0d Hybrid Search '{query[:40]}': {len(results)} results in {elapsed_ms}ms (RRF)"
    )
    client.close()

    return {
        "query_vector": query_vector,
        "results": results,
        "cached": False,
        "search_time_ms": elapsed_ms,
    }


# ── RAG Chain (LangChain LCEL) ─────────────────────────────────────────────────

PROMPT_TEMPLATE = """Bạn là Trợ lý Số Cán bộ Xã chuyên trách về nghiệp vụ và thủ tục hành chính cấp cơ sở. Nhiệm vụ của bạn là giải đáp thông tin cho cán bộ, nhân dân dựa TRÊN TÀI LIỆU THAM KHẢO được cung cấp bên dưới.

🚨 QUY TẮC NGHIÊM NGẶT:
1. CHỈ sử dụng thông tin có trong "TÀI LIỆU THAM KHẢO". Nếu thông tin không có, hãy lịch sự từ chối trả lời và nói rằng tài liệu hiện tại không đề cập đến.
2. KHÔNG tự bịa đặt, suy đoán quy trình hoặc các văn bản pháp luật không có trong nguồn.
3. LUÔN trích dẫn nguồn ([số thứ tự nguồn]) ngay sau thông tin bạn lấy từ tài liệu.
4. Trình bày thông tin chuyên nghiệp, rõ ràng bằng Markdown.
5. Giữ thái độ trung thực, chuẩn mực của một cán bộ hành chính.

TÀI LIỆU THAM KHẢO:
{context}

LỊCH SỬ HỘI THOẠI:
{history}

CÂU HỎI CỦA NGƯỜI DÙNG: {question}

Hãy đưa ra câu trả lời chuyên nghiệp:"""


def build_context_from_results(results: list[dict]) -> str:
    """Xay dung context string va gioi han do dai tung chunk."""
    MAX_CONTEXT_CHARS = 800  # Gioi han an toan de giam bot thoi gian sinh
    parts = []
    for i, r in enumerate(results, 1):
        source = r.get("source", "unknown")
        page = r.get("page_num", "?")
        score = r.get("score", 0)

        # Cat noi dung neu vuot qua gioi han
        content = r.get("content", "").strip()
        if len(content) > MAX_CONTEXT_CHARS:
            content = content[:MAX_CONTEXT_CHARS] + "..."

        parts.append(
            f"[{i}] Nguon: {source} (trang {page}, score: {score:.3f})\n{content}"
        )

    context_str = "\n\n".join(parts)
    logger.info(
        f"Context built: {len(results)} sources, {len(context_str)} characters."
    )
    return context_str


def rag_chat(query: str, history: list[dict] = None, department: str = "public") -> dict:
    """
    Hàm chat chính dùng LangChain LCEL pipeline.
    Returns: {answer, sources, query_vector, search_time_ms, cached}
    """
    if history is None:
        history = []

    # 1. Vector search (với cache)
    search_result = search_vectors(query, department=department)
    results = search_result["results"]

    if not results:
        return {
            "answer": "Tôi chưa tìm được thông tin liên quan trong cơ sở dữ liệu hiện có. Vui lòng upload thêm tài liệu.",
            "sources": [],
            "query_vector": search_result["query_vector"][
                :10
            ],  # Chỉ trả 10 dims đầu để demo
            "query_vector_dim": len(search_result["query_vector"]),
            "search_time_ms": search_result["search_time_ms"],
            "cached": search_result["cached"],
        }

    # 2. Build context
    context = build_context_from_results(results)

    # 3. Build history string
    history_str = ""
    if history:
        recent = history[-4:]
        history_str = "\n".join(
            f"{'Người dùng' if m['role'] == 'user' else 'Bot'}: {m['content']}"
            for m in recent
        )

    # 4. Gọi LLM qua LangChain LCEL
    llm = get_llm()
    prompt = PromptTemplate.from_template(PROMPT_TEMPLATE)

    chain = prompt | llm | StrOutputParser()

    answer = chain.invoke(
        {
            "context": context,
            "history": history_str or "Chưa có lịch sử hội thoại.",
            "question": query,
        }
    )

    return {
        "answer": answer.strip(),
        "sources": results,
        "query_vector": search_result["query_vector"][
            :10
        ],  # 10 dims đầu để hiển thị UI
        "query_vector_dim": len(search_result["query_vector"]),
        "search_time_ms": search_result["search_time_ms"],
        "cached": search_result["cached"],
    }
