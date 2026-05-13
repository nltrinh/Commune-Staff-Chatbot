import io
import logger_config
from typing import List, Dict, Any
from app.core.interfaces import DocumentProcessor

logger = logger_config.get_logger(__name__)

class BasicProcessor(DocumentProcessor):
    """
    Bộ xử lý tài liệu cơ bản, sử dụng các thư viện Python thuần.
    Phù hợp cho môi trường Dev/Local không cài đặt các công cụ nặng.
    """
    def extract_text(self, content: bytes, file_type: str) -> List[Dict[str, Any]]:
        pages = []
        try:
            if file_type == ".txt":
                text = content.decode("utf-8", errors="ignore")
                pages.append({"text": text, "page_num": 1})

            elif file_type == ".pdf":
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(content))
                for i, page in enumerate(reader.pages, 1):
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append({"text": text, "page_num": i})

            elif file_type == ".docx":
                import docx
                doc = docx.Document(io.BytesIO(content))
                full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                if full_text.strip():
                    pages.append({"text": full_text, "page_num": 1})

            elif file_type == ".doc":
                # 1. Thử xem có phải là .docx đổi tên không (rất phổ biến)
                try:
                    import docx
                    doc = docx.Document(io.BytesIO(content))
                    full_text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    if full_text.strip():
                        pages.append({"text": full_text, "page_num": 1})
                        logger.info("Đã trích xuất thành công .doc (thực chất là docx).")
                        return pages
                except Exception:
                    pass

                # 2. Thử dùng striprtf (nếu là định dạng RTF)
                try:
                    # Kiểm tra xem có dấu hiệu RTF không trước khi decode toàn bộ
                    if content.startswith(b"{\\rtf"):
                        from striprtf.striprtf import rtf_to_text
                        text = rtf_to_text(content.decode("utf-8", errors="ignore"))
                        if text.strip():
                            pages.append({"text": text, "page_num": 1})
                            logger.info("Đã trích xuất thành công nội dung .doc qua định dạng RTF.")
                            return pages
                except Exception:
                    pass

                # 3. Nếu là file binary thật sự (Word cũ)
                logger.warning("Không thể đọc tệp .doc binary trên Windows mà không có LibreOffice.")
                msg = (
                    "[CHẾ ĐỘ PHÁT TRIỂN]: Hệ thống nhận diện đây là tệp Word 97-2003 (.doc) nhị phân.\n"
                    "Trên máy Local (Windows), việc trích xuất tự động yêu cầu cài đặt LibreOffice.\n"
                    "LƯU Ý: Khi triển khai lên Vast.ai (Linux), hệ thống sẽ tự động dùng bộ 'ProductionProcessor' "
                    "với đầy đủ công cụ (unstructured + soffice) để đọc tệp này.\n"
                    "Để test ngay bây giờ, vui lòng lưu tệp thành .docx chuẩn."
                )
                pages.append({"text": msg, "page_num": 1})
        except Exception as e:
            logger.error(f"Lỗi trích xuất văn bản ({file_type}): {e}")
            pages.append({"text": f"[LỖI HỆ THỐNG]: {str(e)}", "page_num": 1})
            
        return pages
